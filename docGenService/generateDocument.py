import json
import os
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Length
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pyCaptionTemplates import make_ny_header, make_nj_header, make_fl_header, make_mi_header
from pyGenObjectionTemplates import make_ny_gen_obj, make_nj_gen_obj, make_fl_gen_obj,  make_mi_gen_obj
from pyOutgoingCopy import make_outgoing_instructions
from pyRequestsForProduction import make_requests_for_production


class GenerateBody(object):
    def generate(self, docId):
 
        path = os.path.abspath(os.path.join(os.path.dirname( __file__ ), ".."))
        foundationRogArray = [
            "State your full name, home addresses for the past ten years, your employer for the last 10 years, your current work address, date of birth and social security number.",
            "Identify all insurance carriers or self-insured funds, by name, address, policy numbers and policy limits, for any insurance policy or fund which may provide coverage for any judgment entered against you in this action.",
            "If you contend that any other entity or person, including any other party or the Plaintiff, was responsible for the occurrence and Plaintiff's damages, identify such person(s) or entities, and give a concise statement of the facts upon which you rely in support of your contention.",
            "Identify all experts who will be called to testify at a trial on the merits of this action, the subject matter and substance of such testimony, a summary of the grounds for each opinion; and produce any written report made by the expert concerning those findings and opinions.",
            "Identify any documents and recordings including, but not limited to, pictures, photographs, PowerPoint presentations for use at trial, demonstrative exhibits, computer generated exhibits, electronically stored data, visual aids, overlays, employment records, plats, visual recorded images, audio recordings, cassette tapes, transcripts of testimony, diagrams and objects relative to the occurrence, the scene of the occurrence, Plaintiff's physical condition or statements made by any party or witness. Identify the substance of the item, the date obtained, what is depicted within the item, and the name and address of the present custodian of each item.",
            "If you, your insurance carrier, private investigator, or any other person or entity is in possession of any written, oral or recorded statements by any party or person with personal knowledge relative to the occurrence, indicate the date and time each statement was obtained, the name and address of each person who provided the statement, the contents of the statement, and the name and address of the current custodian of the statement.",
            "Describe any conduct, comment, conversation, statement, or report made by any person at the scene of the occurrence or at any time, concerning fault for the occurrence or facts relevant to any issue in this case. Include in your answer where the conduct, comment, conversation, or statement took place, and in whose presence it was made/observed, as well as the name of the author of such statement, the present custodian of the statement and the address for the custodian.",
            "Since your eighteenth birthday, when you were represented by an attorney or waived the right to be represented by an attorney, state whether you have been found guilty of, or plead guilty to, any crimes other than minor traffic violations (i.e., those traffic offenses without the potential penalty of incarceration) and, if so, state the nature of the offense, the date of each conviction, and the full name of the court where each conviction was entered.",
            "If you are aware of any other case or proceeding involving the incident identified in Plaintiff's Complaint including, but not limited to, civil, criminal or administrative actions, identify the case or action by tribunal, case number, docket number or citation number, the date of any hearing, and indicate any pleas in the case(s) and the disposition of the matter(s).",
        ]

        # Init variables
        reqType = None
        arrLen = None
        respArray = None
        reqArray = None
        reqHeader = None
        respHeader = None
        jsonData = None
        leadAttorneys = None
        clientPosition = None
        respondent = None
        servingParty = None
        firmStreetAddress = None
        firmCity = None
        firmState = None
        firmZip = None
        firmTel = None
        reqFile = None
        respFile = None
        # filePath = f"/Users/kjannette/workspace/ax3/ax3Services/docGenService/Docxstaging/{docId}.docx"
        # f = open(filePath, "rb")
        document = Document()
        dataFile = f"./Docxinfo/{docId}.json"

        # Get case data
        with open(dataFile) as json_data:
            caseInfo = json.load(json_data)

        #caseInfo = jsonData.get("caseInfo")
        
        caption1 = caseInfo["caseCaption1"]
        caption2 = caseInfo["caseCaption2"]
        clientPosition = caseInfo["clientPosition"]
        defendant = caseInfo["defendant"]
        plaintiff = caseInfo["plaintiff"]
        jurisdiction = caseInfo["jurisdiction"]
        reqType = caseInfo["currentRequestType"]
        venue = caseInfo["venue"]
        caseNumber = caseInfo["caseNumber"]
        judge = caseInfo["judge"]

        # Get attorney/firm data
        firm = caseInfo["firm"] 
        leadAttorneys = caseInfo["leadAttorneys"] 
        firmStreetAddress = caseInfo["firmStreetAddress"] 
        firmCity = caseInfo["firmCity"] 
        firmState = caseInfo["state"] 
        firmTel = caseInfo["tel"] 
        firmZip = caseInfo["firmZip"]

        if reqType == "interrogatories":
            #reqFile = f"/var/www/ax3Services/Documents/Requests/Parsedrogs/{docId}/{docId}-jbk-parsedRequests.json"
            #respFile = f"/var/www/ax3Services/Documents/Responses/Rogresp/{docId}/{docId}-jbk-responses.json"
            reqFile = f"{path}/Documents/Requests/interrogatories/{docId}/{docId}-jbk-parsedRequests.json"
            respFile = f"{path}/Documents/Responses/interrogatories/{docId}/{docId}-jbk-responses.json"
        elif reqType == "combined-numbered":
            reqFile = f"{path}/Documents/Requests/combined-numbered/{docId}/{docId}-jbk-parsedRequests.json"
            respFile = f"{path}/Documents/Responses/combined-numbered/{docId}/{docId}-jbk-responses.json"
            #reqFile = f"/var/www/ax3Services/Documents/Requests/combined-numbered/{docId}/{docId}-jbk-parsedRequests.json"
            #respFile = f"/var/www/ax3Services/Documents/Responses/combined-numbered/{docId}/{docId}-jbk-responses.json"
        elif reqType == "production":
            reqFile = f"{path}/Documents/Requests/production/{docId}/{docId}-jbk-parsedRequests.json"
            reqFile = f"{path}/Documents/Requests/production/{docId}/{docId}-jbk-parsedRequests.json"
            #reqFile = f"/var/www/ax3Services/Documents/Requests/production/{docId}/{docId}-jbk-parsedRequests.json"
            #reqFile = f"/var/www/ax3Services/Documents/Requests/production/{docId}/{docId}-jbk-parsedRequests.json"
        elif reqType == "admissions":
            reqFile = f"{path}/Documents/Requests/admissions/{docId}/{docId}-jbk-parsedRequests.json"
            respFile = f"{path}/Documents/Responses/admissions/{docId}/{docId}-jbk-responses.json"
            #reqFile = f"/var/www/ax3Services/Documents/Requests/admissions/{docId}/{docId}-jbk-parsedRequests.json"
            #respFile = f"/var/www/ax3Services/Documents/Responses/admissions/{docId}/{docId}-jbk-responses.json"
        elif reqType == "interrogatories-out":
            reqFile = f"{path}/Documents/RequestsOut/{docId}/{docId}-jbk-requests-out.json"
            #reqFile = f"/var/www/ax3Services/Documents/RequestsOut/{docId}/{docId}-jbk-requests-out.json"

        if reqType == 'interrogatories-out':
            if clientPosition == "Plaintiff":
                respondent = defendant
            elif clientPosition == "Defendant":
                respondent = plaintiff
        else: 
            if clientPosition == "Plaintiff":
                respondent = plaintiff
            elif clientPosition == "Defendant":
                respondent = defendant

        if reqType == 'interrogatories-out':
            if clientPosition == "Plaintiff":
                servingParty = plaintiff
            elif clientPosition == "Defendant":
                servingParty = defendant
        else:
            if clientPosition == "Plaintiff":
                servingParty = caption2
            elif clientPosition == "Defendant":
                servingParty = caption1

        if reqType == 'interrogatories-out':
            if firmState == "mi":
                comesNowString = f"COMES NOW, {clientPosition}, {servingParty}, through counsel, and hereby propounds these Interrogatories and Requests for Production upon {respondent}, to be answered under oath, in writing."
            elif firmState == "ny":
                comesNowString = f"COMES NOW, {clientPosition}, {servingParty}, through counsel, and hereby propounds these Interrogatories and Requests for Production upon {respondent}, to be answered under oath, in writing, in accordance with NY CPLR 3120 - 3130."
            elif firmState == "nj":
                comesNowString = f"Comes now ${clientPosition}, ${servingParty}, through counsel, and hereby propounds these Interrogatories and Requests for Production upon ${respondent}, to be answered to be answered under oath, in writing, in accordance with NJ R. 4:17-1 - 4:18-1, + et. seq. All questions must be answered unless the court otherwise orders or unless a claim of privilege or protective order is made in accordance with R. 4:17-1(b)(3)."
            elif firmState == "fl":
                comesNowString = f"COMES NOW,  ${clientPosition}, ${servingParty}, through counsel, and hereby propounds these Interrogatories and Requests for Production upon ${respondent}, to be answered to be answered under oath, in writing, in accordance with Rule 1.340, Florida Rules of Civil Procedure."
        else: 
            comesNowString = f"COMES NOW, Respondent(s), {respondent}, through counsel, in response to the {reqType} served by {servingParty}, and states as follows:"

        if reqType == "interrogatories":
            reqHeader = "INTERROGATORY NO."
            respHeader = "RESPONSE TO INTERROGATORY NO."
            mainHeader = "RESPONSE TO INTERROGATORIES"
            mainHeader2 = "RESPONSES"
        elif reqType == "admissions":
            reqHeader = "Request for Admission No."
            respHeader = "Response to Request for Admission  No."
            mainHeader = "RESPONSE TO REQUEST FOR ADMISSIONS"
            mainHeader2 = "RESPONSES"
        elif reqType == "production":
            reqHeader = "Request for Production No."
            respHeader = "Response to Request for Production  No."
            mainHeader = "RESPONSE TO REQUEST FOR PRODUCTON"
            mainHeader2 = "RESPONSES"
        elif reqType == 'interrogatories-out':
            reqHeader = "Request no."
            respHeader = "Response to Request for Production  No."
            mainHeader = "INTERROGATORIES AND REQUEST FOR PRODUCTION OF DOCUMENTS"
            mainHeader2 = "INTERROGATORIES"
        else:
            reqHeader = "Request No."
            respHeader = "Response to Request No."
            mainHeader = "RESPONSE TO INTERROGATORIES AND REQUEST FOR PRODUCTION"
            mainHeader2 = "RESPONSES"

        # Create header
        if firmState == "mi":
            document = make_mi_header(
                document,
                jurisdiction,
                venue,
                caption1,
                caption2,
                mainHeader,
                caseNumber,
                judge,
                clientPosition,
            )
        elif firmState == "ny":
            document = make_ny_header(
                document,
                jurisdiction,
                venue,
                caption1,
                caption2,
                mainHeader,
                caseNumber,
                judge,
                clientPosition,
            )
        elif firmState == "nj":
            document = make_nj_header(
                comesNowString,
                firm,
                leadAttorneys,
                document,
                jurisdiction,
                venue,
                caption1,
                caption2,
                mainHeader,
                caseNumber,
                judge,
                firmStreetAddress,
                firmCity,
                firmState,
                firmTel,
                firmZip,
                clientPosition,
            )
        elif firmState == "fl":
            document = make_fl_header(
                comesNowString,
                firm,
                leadAttorneys,
                document,
                jurisdiction,
                venue,
                caption1,
                caption2,
                mainHeader,
                caseNumber,
                judge,
                firmStreetAddress,
                firmCity,
                firmState,
                firmTel,
                firmZip,
                clientPosition,
            )
        elif firmState == None:
            document = make_ny_header(
                document,
                jurisdiction,
                venue,
                caption1,
                caption2,
                mainHeader,
                caseNumber,
                judge,
                clientPosition,
            )

        paragraph = document.add_paragraph(f"{comesNowString}")
        paragraph.paragraph_format.line_spacing = Pt(24)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Add Instructions/Definitions if outgoing, else add General Objections 
        if reqType == "interrogatories-out":
            document = make_outgoing_instructions(document, clientPosition, servingParty)
        else:
            if firmState == "mi":
                document = make_mi_gen_obj(document, clientPosition, servingParty)
            elif firmState == "ny":
                document = make_ny_gen_obj(document, clientPosition, servingParty)
            elif firmState == "nj":
                document = make_nj_gen_obj(document, clientPosition, servingParty)
            elif firmState == "fl":
                document = make_fl_gen_obj(document, clientPosition, servingParty)

        # Add main heading two
        p = document.add_paragraph()
        p.add_run(f"{mainHeader2}").underline = True
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # INIT responses to iterate if doc is responsive, not outgoing
        if (reqType != "interrogatories-out"):
            json_resp_data = open(respFile)
            respData = json.load(json_resp_data)
            respArray = respData[0]["responses"]
            arrLen = len(respArray)

        # INIT requests and begin iteration to generate copy
        if (reqType == "interrogatories-out"):
            with open(reqFile) as json_data:
                dataz = json.load(json_data)
                reqArray = dataz[0]["requests"]
                filtered = []
                for req in reqArray:
                    filtered.append(req["text"])
                
                joinedReqArray = [*foundationRogArray, *filtered]
                arrLen2 = len(joinedReqArray)
                count = 1
            
                for req in joinedReqArray:
                    if count == arrLen2:
                        break
                    
                    paragraph = document.add_paragraph(f"{count}. {req}")
                    paragraph.paragraph_format.line_spacing = Pt(20)
                    paragraph.paragraph_format.space_after = Pt(12)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    count = count + 1

            document = make_requests_for_production(document, clientPosition, servingParty)

        elif (reqType != "interrogatories-out"):
            with open(reqFile) as json_data:
                dataz = json.load(json_data)
                reqArray = dataz[0]["requests"]
                count = 1

                for req in reqArray:
                    if count == arrLen:
                        break
                    text = req["text"]
                    paragraph = document.add_paragraph(f"{reqHeader} {count}:")
                    paragraph = document.add_paragraph(text)
                    paragraph.paragraph_format.line_spacing = Pt(20)
                    paragraph.paragraph_format.space_after = Pt(12)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragraph = document.add_paragraph(f"{respHeader} {count}:")
                    paragraph.paragraph_format.space_before = Pt(24)
                    paragraph = document.add_paragraph(respArray[count]["text"])
                    paragraph.paragraph_format.line_spacing = Pt(24)
                    paragraph.paragraph_format.space_after = Pt(12)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    count = count + 1

        paragraph = document.add_paragraph(
            f"For the {clientPosition}, attorney {leadAttorneys} on this ______ day of _______________________, 2024."
        )
        paragraph.paragraph_format.space_before = Pt(24)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph = document.add_paragraph(
            "                                                                                                            ___________________________________"
        )
        paragraph.paragraph_format.space_before = Pt(24)
        paragraph.paragraph_format.space_after = Pt(12)
        document.save(f"{path}/Docxfinal/{docId}.docx")
        
        respBody = json.dumps({'Status': 'Success'})
        return respBody

# document.save(
#   f"/Users/kjannette/workspace/ax3/ax3Services/Docxfinal/{docId}.docx"
#         document.save(f"/var/www/ax3Services/Docxfinal/{docId}.docx")
# Uncomment for development/smoke testing
#genBod = GenerateBody()

#genBod.generate("eb75d30a-d58a-4b2e-80ba-695c8a79a1e6")
# reqFile = f"/Users/kjannette/workspace/ax3/ax3Services/Documents/Requests/Parsedrogs/{docId}/{docId}-jbk-parsedRequests.json"
# respFile = f"/Users/kjannette/workspace/ax3/ax3Services/Documents/Responses/Rogresp/{docId}/{docId}-jbk-responses.json"

# reqFile = f"/Users/kjannette/workspace/ax3/ax3Services/Documents/Requests/Parsedrogs/{docId}/{docId}-jbk-parsedRequests.json"
# respFile = f"/Users/kjannette/workspace/ax3/ax3Services/Documents/Responses/Rogresp/{docId}/{docId}-jbk-responses.json"
