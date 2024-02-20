import json
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Length
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pyCaptionTemplates import make_ny_header, make_nj_header, make_fl_header
from pyGenObjectionTemplates import make_ny_gen_obj, make_nj_gen_obj, make_fl_gen_obj


class GenerateBody(object):
    def generate(self, docId):
        # Init variables
        reqType = None
        arrLen = None
        respArray = None
        reqArray = None
        reqHeader = None
        respHeader = None
        jsonData = None
        leadAttorneys = None
        clientPosition = None
        respondent = None
        servingParty = None
        firmStreetAddress = None
        firmCity = None
        firmState = None
        firmZip = None
        firmTel = None
        # filePath = f"/Users/kjannette/workspace/ax3/ax3Services/docGenService/Docxstaging/{docId}.docx"
        # f = open(filePath, "rb")
        document = Document()

        # reqFile = f"/Users/kjannette/workspace/ax3/ax3Services/Documents/Requests/Parsedrogs/{docId}/{docId}-jbk-parsedRequests.json"
        # respFile = f"/Users/kjannette/workspace/ax3/ax3Services/Documents/Responses/Rogresp/{docId}/{docId}-jbk-responses.json"

        reqFile = f"/var/www/ax3Services/Documents/Requests/Parsedrogs/{docId}/{docId}-jbk-parsedRequests.json"
        respFile = f"/var/www/ax3Services/Documents/Responses/Rogresp/{docId}/{docId}-jbk-responses.json"

        dataFile = f"Docxinfo/{docId}.json"

        # Get case data
        with open(dataFile) as json_data:
            jsonData = json.load(json_data)

        caption1 = jsonData.get("caseCaption1")
        caption2 = jsonData.get("caseCaption2")
        clientPosition = jsonData.get("clientPosition")
        defendant = jsonData.get("defendant")
        plaintiff = jsonData.get("defendant")
        jurisdiction = jsonData.get("jurisdiction")
        reqType = jsonData.get("currentRequestType")
        venue = jsonData.get("venue")
        caseNumber = jsonData.get("caseNumber")
        judge = jsonData.get("judge")

        # Get attorney/firm data
        firm = jsonData.get("firm")
        leadAttorneys = jsonData.get("leadAttorneys")
        firmStreetAddress = jsonData.get("firmStreetAddress")
        firmCity = jsonData.get("firmCity")
        firmState = jsonData.get("state")
        firmTel = jsonData.get("tel")
        firmZip = jsonData.get("firmZip")
        print(
            "````````````````````````````````````````````````````````firmState",
            firmState,
        )
        if clientPosition == "Plaintiff":
            respondent = plaintiff
        elif clientPosition == "Defendant":
            respondent = defendant

        if clientPosition == "Plaintiff":
            servingParty = caption2
        elif clientPosition == "Defendant":
            servingParty = caption1

        comesNowString = f"COMES NOW, Respondent(s), {respondent}, through counsel, in response to the {reqType} served by {servingParty}, and states as follows:"

        if reqType == "interrogatories":
            reqHeader = "INTERROGATORY NO."
            respHeader = "RESPONSE TO INTERROGATORY NO."
            mainHeader = "RESPONSE TO INTERROGATORIES"
            mainHeader2 = "RESPONSES"
        elif reqType == "admissions":
            reqHeader = "Request for Admission No."
            respHeader = "Response to Request for Admission  No."
            mainHeader = "RESPONSE TO REQUEST FOR ADMISSIONS"
            mainHeader2 = "RESPONSES"
        elif reqType == "production":
            reqHeader = "Request for Production No."
            respHeader = "Response to Request for Production  No."
            mainHeader = "RESPONSE TO REQUEST FOR PRODUCTON"
            mainHeader2 = "RESPONSES"
        else:
            reqHeader = "Request No."
            respHeader = "Response to Request No."
            mainHeader = "RESPONSE TO INTERROGATORIES AND REQUEST FOR PRODUCTION"
            mainHeader2 = "RESPONSES"

        # Create header
        if firmState == "ny":
            document = make_ny_header(
                document,
                jurisdiction,
                venue,
                caption1,
                caption2,
                mainHeader,
                caseNumber,
                judge,
                clientPosition,
            )
        elif firmState == "nj":
            document = make_nj_header(
                comesNowString,
                firm,
                leadAttorneys,
                document,
                jurisdiction,
                venue,
                caption1,
                caption2,
                mainHeader,
                caseNumber,
                judge,
                firmStreetAddress,
                firmCity,
                firmState,
                firmTel,
                firmZip,
                clientPosition,
            )
        elif firmState == "fl":
            document = make_fl_header(
                comesNowString,
                firm,
                leadAttorneys,
                document,
                jurisdiction,
                venue,
                caption1,
                caption2,
                mainHeader,
                caseNumber,
                judge,
                firmStreetAddress,
                firmCity,
                firmState,
                firmTel,
                firmZip,
                clientPosition,
            )
        elif firmState == None:
            document = make_ny_header(
                document,
                jurisdiction,
                venue,
                caption1,
                caption2,
                mainHeader,
                caseNumber,
                judge,
                clientPosition,
            )

        paragraph = document.add_paragraph(f"{comesNowString}")
        paragraph.paragraph_format.line_spacing = Pt(24)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Add General Objections
        if firmState == "New York":
            document = make_ny_gen_obj(document, clientPosition, servingParty)
        elif firmState == "New Jersey":
            document = make_nj_gen_obj(document, clientPosition, servingParty)
        elif firmState == "Florida":
            document = make_fl_gen_obj(document, clientPosition, servingParty)

        # Add main heading two
        p = document.add_paragraph()
        p.add_run(f"{mainHeader2}").underline = True
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # INIT responses to iterate
        json_resp_data = open(respFile)
        respData = json.load(json_resp_data)
        respArray = respData[0]["responses"]
        arrLen = len(respArray)

        # INIT requests and begin iteration to generate copy
        with open(reqFile) as json_data:
            dataz = json.load(json_data)
            reqArray = dataz[0]["requests"]
            count = 1

            for req in reqArray:
                if count == arrLen:
                    break
                text = req["text"]
                paragraph = document.add_paragraph(f"{reqHeader} {count}:")
                paragraph = document.add_paragraph(text)
                paragraph.paragraph_format.line_spacing = Pt(20)
                paragraph.paragraph_format.space_after = Pt(12)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph = document.add_paragraph(f"{respHeader} {count}:")
                paragraph.paragraph_format.space_before = Pt(24)
                paragraph = document.add_paragraph(respArray[count]["text"])
                paragraph.paragraph_format.line_spacing = Pt(24)
                paragraph.paragraph_format.space_after = Pt(12)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                count = count + 1

        paragraph = document.add_paragraph(
            f"For the {clientPosition}, attorney {leadAttorneys} on this ______ day of _______________________, 2024."
        )
        paragraph.paragraph_format.space_before = Pt(24)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph = document.add_paragraph(
            "                                                                                                            ___________________________________"
        )
        paragraph.paragraph_format.space_before = Pt(24)
        paragraph.paragraph_format.space_after = Pt(12)
        document.save(f"/var/www/ax3Services/Docxfinal/{docId}.docx")


# document.save(
#   f"/Users/kjannette/workspace/ax3/ax3Services/Docxfinal/{docId}.docx"
# )
# Uncomment for development/smoke testing
# genBod = GenerateBody()
# genBod.generate("a93d9b76-28aa-4b8c-83e4-e87ec6294c4b")
