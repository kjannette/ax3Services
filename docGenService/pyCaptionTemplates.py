import json
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Length
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# New York *******************************************************************/
def make_ny_header(
    document,
    jurisdiction,
    venue,
    caption1,
    caption2,
    mainHeader,
    caseNumber,
    judge,
    clientPosition,
):
    print(
        "______________________________________________________________NEW YORK MAKE HEADER FIRED"
    )
    paragraph = document.add_paragraph(f"{jurisdiction}")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph = document.add_paragraph(f"{venue}")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph = document.add_paragraph(
        "-------------------------------------------------------------X"
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph = document.add_paragraph(f"{caption1}")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph = document.add_paragraph(
        f"                                                Plaintiff(s)                                                                  Index No.: {caseNumber}"
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph = document.add_paragraph(
        f"- against -                                                                                                                    Judge: {judge}"
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(16)
    paragraph = document.add_paragraph(f"{caption2}")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph = document.add_paragraph(
        "                                                Defendant(s)"
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph = document.add_paragraph(
        "-------------------------------------------------------------X"
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)

    p = document.add_paragraph()
    p.add_run(f"{clientPosition.upper()}'S {mainHeader}").underline = True
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return document


# Florida *******************************************************************/


def make_fl_header(
    comesNowString,
    firm,
    leadAttorneys,
    document,
    jurisdiction,
    venue,
    caption1,
    caption2,
    mainHeader,
    caseNumber,
    judge,
    firmStreetAddress,
    firmCity,
    firmState,
    firmTel,
    firmZip,
    clientPosition,
):
    print(
        "______________________________________________________________FLORIDA MAKE HEADER FIRED"
    )
    jurisdictionUpper = jurisdiction.upper()
    venueUpper = venue.upper()

    p = document.add_paragraph()
    p.add_run(f" IN THE CIRCUIT COURT OF THE {jurisdictionUpper}")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = document.add_paragraph()
    p.add_run(f"OF FLORIDA, IN AND FOR {venueUpper} COUNTY")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph = document.add_paragraph(f"{caption1}")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph = document.add_paragraph(
        f"                                                Plaintiff(s)                                                            Case No.: {caseNumber}"
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph = document.add_paragraph(f"v.")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(16)
    paragraph = document.add_paragraph(f"{caption2}")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph = document.add_paragraph(
        "                                                Defendant(s)"
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph = document.add_paragraph(
        "________________________________________________________________________/"
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)

    p = document.add_paragraph()
    p.add_run(f"{clientPosition.upper()}'S {mainHeader}").underline = True
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return document


# New Jersey *******************************************************************/
def make_nj_header(
    comesNowString,
    firm,
    leadAttorneys,
    document,
    jurisdiction,
    venue,
    caption1,
    caption2,
    mainHeader,
    caseNumber,
    judge,
    firmStreetAddress,
    firmCity,
    firmState,
    firmTel,
    firmZip,
    clientPosition,
):
    paragraph = document.add_paragraph(f"{firm}")
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph = document.add_paragraph(f"By: {leadAttorneys}")
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph = document.add_paragraph(f"{firmStreetAddress}")
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph = document.add_paragraph(f"{firmCity}, {firmState} {firmZip}")
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph = document.add_paragraph(f"Tel: {firmTel}")
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph = document.add_paragraph(
        f"_______________________________________________________________"
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph = document.add_paragraph(f"{caption1:<92} ")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph = document.add_paragraph(
        f"                                                                                                                  {jurisdiction}          "
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph = document.add_paragraph(
        f"                                                                                                                  {venue}          "
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph = document.add_paragraph(
        f"                      Plaintiff                                                                          "
    )
    paragraph = document.add_paragraph(
        f"                                                                                                                           "
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph = document.add_paragraph(
        f"                      v.                                                                                           Docket No,{caseNumber}"
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph = document.add_paragraph(
        f"                                                                                                                         "
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph = document.add_paragraph(
        f"                                                                                                                          "
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph = document.add_paragraph((f"{caption2:<90}  "))

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph = document.add_paragraph(
        f"                                                                                                        "
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph = document.add_paragraph(
        f"                      Defendant                                                                      "
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph = document.add_paragraph(
        f"                                                                                                              "
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph = document.add_paragraph(
        f"_______________________________________________________________"
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    return document
