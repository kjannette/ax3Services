import json
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Length
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Michigan *****************************************************************/
def make_mi_gen_obj(document, clientPosition, servingParty):
    objectionsArray = [
        f"All of {clientPosition}'s responses herein are subject to the following objections, in addition to any and all specific objections set forth in the responses to individual requests.",
        f"In responding to these requests, {clientPosition} does not admit or concede any presumptions or assumptions made in the propounding parties' definitions, or in the requests themselves.",
        f"{clientPosition} objects to any and all requests to the extent they seek or may be interpreted to seek disclosure of information not within the scope of MCR 2.302(B), or not within the scope of what is permitted under any applicable Preliminary Conference or Case Scheduling Order entered in this case, and {clientPosition} reserves all rights to contest any such matters in any other context or proceeding where they may be relevant.",
        f"{clientPosition} objects to all requests that seek disclosure of information which (a) is subject to attorney-client privilege, or (b) the 'work product' doctrine; (c) is subject to the required reports privilege; (d) is subject to a joint defense or common interest privilege; (e) was generated in anticipation of litigation or for trial, (f) relates to the identity or opinions of experts who have been retained or employed in anticipation of litigation and who are not expected to be called as witnesses at trial; (g) is protected as a trade secret; (h) is subject to a protective order or confidentiality order or agreement which was entered or made in another matter, and/or (i) is otherwise privileged, protected from disclosure, or beyond the scope of discovery under applicable rules and laws.  Further, {clientPosition} does not intend to disclose or produce any such information,  and any disclosure of such  information is inadvertent, and all rights to demand return and/or destruction of any such information are reserved.",
        f"{clientPosition}  objects to the propounding parties' requests insofar as they seek a proposition of law and/or the formulation of a legal theory, or seek contentions regarding factual matters as to which essential discovery is incomplete. {clientPosition}'s current responses to such requests necessarily cannot present all information {clientPosition} may ultimately discover and utilize or rely upon in this matter. {clientPosition} thus reserves all rights to supplement or amend its responses in accordance with applicable rules, laws, orders or agreements of the parties, if and when circumstances warrant.",
    ]
    print(
        "______________________________________________________________NEW YORK MAKE OBJECTION FIRED"
    )
    p = document.add_paragraph()
    p.add_run("GENERAL OBJECTIONS").underline = True
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    arrLen = len(objectionsArray)
    count = 0

    for obj in objectionsArray:
        if count == arrLen:
            break

        paragraph = document.add_paragraph(f"{count + 1}. {obj}")
        paragraph.paragraph_format.line_spacing = Pt(20)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        count = count + 1

    return document

# New York *****************************************************************/
def make_ny_gen_obj(document, clientPosition, servingParty):
    objectionsArray = [
        f"All of {clientPosition}'s responses herein are subject to the following objections, in addition to any and all specific objections set forth in the responses to individual requests.",
        f"In responding to these requests, {clientPosition} does not admit or concede any presumptions or assumptions made in the propounding parties' definitions, or in the requests themselves.",
        f"{clientPosition} objects to any and all requests to the extent they seek or may be interpreted to seek disclosure of information not within the scope of New York CPLR 3120, 3130 - 33, or not within the scope of what is permitted under any applicable Preliminary Conference or Case Scheduling Order entered in this case, and {clientPosition} reserves all rights to contest any such matters in any other context or proceeding where they may be relevant.",
        f"{clientPosition} objects to all requests that seek disclosure of information which (a) is subject to attorney-client privilege, or (b) the 'work product' doctrine; (c) is subject to the self-critical analysis privilege; (d) is subject to the required reports privilege; (e) is subject to a joint defense or common interest privilege; (f) was generated in anticipation of litigation or for trial, (g) relates to the identity or opinions of experts who have been retained or employed in anticipation of litigation and who are not expected to be called as witnesses at trial; (h) is protected as a trade secret; (i) is subject to a protective order or confidentiality order or agreement which was entered or made in another matter, and/or (j) is otherwise privileged, protected from disclosure, or beyond the scope of discovery under applicable rules and laws.  Further, {clientPosition} does not intend to disclose or produce any such information,  and any disclosure of such  information is inadvertent, and all rights to demand return and/or destruction of any such information are reserved.",
        f"{clientPosition}  objects to the propounding parties' requests insofar as they seek a proposition of law and/or the formulation of a legal theory, or seek contentions regarding factual matters as to which essential discovery is incomplete. {clientPosition}'s current responses to such requests necessarily cannot present all information {clientPosition} may ultimately discover and utilize or rely upon in this matter. {clientPosition} thus reserves all rights to supplement or amend its responses in accordance with applicable rules, laws, orders or agreements of the parties, if, as and when circumstances may warrant.",
    ]
    print(
        "______________________________________________________________NEW YORK MAKE OBJECTION FIRED"
    )
    p = document.add_paragraph()
    p.add_run("GENERAL OBJECTIONS").underline = True
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    arrLen = len(objectionsArray)
    count = 0

    for obj in objectionsArray:
        if count == arrLen:
            break

        paragraph = document.add_paragraph(f"{count + 1}. {obj}")
        paragraph.paragraph_format.line_spacing = Pt(20)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        count = count + 1

    return document


# New Jersey *******************************************************************/
def make_nj_gen_obj(document, clientPosition, servingParty):
    objectionsArray = [
        f"All of {clientPosition}’s responses to the discovery requests being answered are subject to the following objections, in addition to any and all objections stated in the answers to each individual request.",
        f"""{clientPosition} objects to the definitions and instructions included in {servingParty}'s discovery requests, to the extent that: (a) the definitions or instructions are inconsistent with any applicable statutes, regulations, laws, legal precedents, or the terms of any applicable agreements or other legal documents; (b) the definitions or instructions seek to impose on {clientPosition} obligations that exceed the requirements of the New Jersey Rules of Court; and/or, (c)  the definitions are overly broad or inclusive, and presume or assume unproven assertions of fact or law.""",
        f"In responding to these requests, {clientPosition} does not admit or concede any presumptions or assumptions embedded in {servingParty}'s definitions.",
        f"{clientPosition} objects to any and all requests to the extent they seek or may be interpreted to seek disclosure of information not within the scope of R. 4:10-2(a) or not within the scope of what is permitted under any applicable Case Management Order entered in this case, and {clientPosition} reserves all rights to contest any such matters in any other context or proceeding where they may be relevant.",
        f"{clientPosition} objects to any and all requests to the extent they seek or may be interpreted to seek disclosure of any information which (a) is subject to the attorney-client privilege; (b) is covered by the 'work product' doctrine; (c) is subject to the self-critical analysis privilege; (d) is subject to the required reports privilege; (e) is subject to a joint defense or common interest privilege; (f) was generated in anticipation of litigation or for trial by or for {clientPosition} or any representatives of {clientPosition} including attorneys, consultants or agents; (g) relates to the identity or opinions of consultants or experts who have been retained or specially employed in anticipation of litigation and who are not expected to be called as witnesses at trial; (h) is protected as a trade secret; (i) is subject to a protective order or confidentiality order or agreement which was entered or made in another matter, to the extent the same prevents disclosure in this matter; and/or (j) is otherwise privileged, protected from disclosure, or beyond the scope of discovery under applicable rules and laws. {clientPosition} does not intend to disclose or produce any such information in response to the request being answered, and the following responses should be read accordingly. Any disclosure of information which is privileged or otherwise protected from disclosure is inadvertent, and all rights to demand return and/or destruction of any such information are reserved.",
        f"{clientPosition} objects to the propounding parties' requests insofar as they seek a proposition of law and/or the formulation of a legal theory, or seek contentions regarding factual matters as to which essential discovery is incomplete. {clientPosition}'s current responses to such requests necessarily cannot present all information {clientPosition} may ultimately discover and utilize or rely upon in this matter. {clientPosition} thus reserves all rights to supplement or amend its responses in accordance with applicable rules, laws, orders or agreements of the parties as circumstances may warrant.",
    ]
    print(
        "______________________________________________________________NEW JERSEY MAKE OBJECTION FIRED"
    )
    p = document.add_paragraph()
    p.add_run("GENERAL OBJECTIONS").underline = True
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    arrLen = len(objectionsArray)
    count = 0

    for obj in objectionsArray:
        if count == arrLen:
            break

        paragraph = document.add_paragraph(f"{count + 1}. {obj}")
        paragraph.paragraph_format.line_spacing = Pt(20)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        count = count + 1

    return document


# Florida *******************************************************************/
def make_fl_gen_obj(document, clientPosition, servingParty):
    objectionsArray = [
        f"All of {clientPosition}’s responses to the discovery requests answered herein are subject to the following objections, in addition to any and all objections stated in the answers to each individual request.",
        f"""{clientPosition} objects to the definitions and instructions included in {servingParty}'s discovery requests, to the extent that: (a) the definitions or instructions are inconsistent with any applicable statutes, rules, legal precedents, or the terms of any applicable agreements or orders; (b) the definitions or instructions seek to impose on {clientPosition} obligations that exceed the requirements of the Florida Rules of Civil Procedure; and/or, (c)  the definitions are overly broad or inclusive, and presume or assume unproven assertions of fact or law.""",
        f"In responding to these requests, {clientPosition} does not admit or concede any presumptions or assumptions embedded in {servingParty}'s definitions.",
        f"{clientPosition} objects to any and all requests to the extent they seek or may be interpreted to seek disclosure of information not within the scope of what is permitted under any applicable Case Management Order entered in this case, and {clientPosition} reserves all rights to contest any such matters in any other context or proceeding where they may be relevant.",
        f"{clientPosition} objects to any requests to the extent they seek disclosure of information which is subject to (a) the attorney-client privilege; (b) the 'work product' doctrine; (c) the spousal communications privilege; (e) a joint defense or common interest privilege; or (f) was generated in anticipation of litigation by or for {clientPosition} or any representatives of {clientPosition}; or (g) that relates to the identity or opinions of experts retained in anticipation of litigation and who are not expected to be called as trial witnesses; or (h) is protected as a trade secret; (i) subject to a protective order in another matter; or (j) is otherwise privileged, protected from disclosure, or beyond the scope of discovery. {clientPosition} does not intend to disclose any such information, and disclosure of any information which is privileged or protected from disclosure is inadvertent. All rights to demand return or destruction of such information are reserved.",
        f"{clientPosition} objects to the propounding parties' requests insofar as they seek a conclusion of law and/or formulation of a legal theory, or seek contentions regarding factual matters as to which essential discovery is incomplete. {clientPosition}'s current responses to such requests necessarily cannot present all information {clientPosition} may ultimately discover and utilize or rely upon in this matter. {clientPosition} thus reserves all rights to supplement or amend its responses in accordance with applicable rules, laws, orders or agreements of the parties as circumstances may warrant.",
    ]
    print(
        "______________________________________________________________FLORIDA MAKE OBJECTION FIRED"
    )
    
    p = document.add_paragraph()
    p.add_run("GENERAL OBJECTIONS").underline = True
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    arrLen = len(objectionsArray)
    count = 0

    for obj in objectionsArray:
        if count == arrLen:
            break

        paragraph = document.add_paragraph(f"{count + 1}. {obj}")
        paragraph.paragraph_format.line_spacing = Pt(20)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        count = count + 1

    return document
