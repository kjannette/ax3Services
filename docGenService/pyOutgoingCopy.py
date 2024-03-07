import json
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Length
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def make_outgoing_instructions(document, clientPosition, servingParty):
    print(
        "______________________________________________________________OUTGOING FIRED"
    )
    instructionsArray = [
      "These requests are continuing in nature and require you to file supplementary or amended responses if you obtain further, different or contradictory information at any time prior to a trial of this action.",
      "Unless otherwise stated, these requests refer to the time, place, and circumstances of the occurrences described and/or complained of in the Complaint and which form the basis for this action.",
      "Where name and identity of a person is required, state the full name, home address and business address, if known.",
      "When the identity of a corporation or other business entity is requested, state the entity's principal place of business, address and telephone number.  Also, set forth the individual or individuals with knowledge of the facts and circumstances that form the basis for this action.",
      "Where knowledge or information in possession of a party is requested, such request includes knowledge of the party's agents, representatives, and, unless privileged,  attorneys. ",
      "Where the respondent is a corporation or other business entity, state the name, address and title of persons supplying the information and making the affidavit, and state with particularity the source of his or her information and basis of knowledge.",
      "If you are unable to provide an accurate answer to an interrogatory, or are uncertain regarding a response to any item, please provide your best estimate of the information requested, specifying that the answer is merely an estimate, and that an accurate response cannot be given by you to the item.",
      "If you withhold any information or documents on the basis of privilege, confidentiality, relevance, or for any other reason whatsoever  please identify it, providing sufficient information to identify the withheld documents or items for this party and the Court to determine whether the documents, items or information should be disclosed.",
    ]

    definitionsArray = [
      "'You' refers to the party to whom the requests are addressed and any additional parties or entities as described in the foregoing instructions and these definitions.",
      "“Person”, unless otherwise specified”, includes the plural as well as the singular, and includes in its meaning any natural person, or artificial or legal entity, including corporations, partnerships, joint ventures, associations, governmental agencies, groups, organizations, and any and every other form of entity cognizable at law.",
      "“Document” means and includes all written and graphic matter of every kind and description, whether printed or produced by any process or by hand, whether final draft or reproduction, whether in the actual or constructive possession, custody or control of the respondent to a given interrogatory, including any and all written letters, correspondence, memoranda, notes, statements, transcripts, files, charters, articles of incorporation, securities, bonds, stocks, certificates of deposit, evidences of debt, contracts, agreements, licenses, memoranda or notes of telephone or personal conversations, work papers, tapes, charts, reports, books, ledgers, telegrams, sound recordings, books of account, customer account statements, financial statements, catalogs, checks, check stubs, and written statements of witnesses or other persons having knowledge pertaining to the pertinent facts requested or relating to the interrogatory or subpart thereof, whether or not these documents are claimed to be privileged against disclosure.",
      "“Identify” means, when used with reference to an individual person, organization, corporation or association, to state the full name, home and work addresses, e-mail address, home and business telephone numbers, present or last known position and business affiliation and position both in the past and at the time said interrogatory or subpart thereof is being responded to.",
      "The term “identify” further means, when used with reference to a document, to state the date it was created, its author, the signatory, if different, the addressee, the recipient of all copies, the type of document (e.g. chart, memorandum, letter, or other written document) and its present or last known custodian.",
      "“Describe” means that the person or entity to whom the interrogatory or subpart thereof is directed should state what is requested to be described, including all facts and opinions known and held regarding, relating to, or pertinent to what is requested to be described, and (i) the identity of each person or entity involved or having any knowledge of each fact or opinion that relates to what is so described, (ii) the identity of each document evidencing the answer or response given or relating, referring or pertaining to said subject-matter in any way, and (iii) all relevant or material dates and time periods, specifying the way in which said dates or time periods are pertinent to the subject-matter described.",
    ]

    p = document.add_paragraph()
    p.add_run("INSTRUCTIONS").underline = True
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    arrLen = len(instructionsArray)
    count = 0

    for inst in instructionsArray:
        if count == arrLen:
            break

        paragraph = document.add_paragraph(f"{count + 1}. {inst}")
        paragraph.paragraph_format.line_spacing = Pt(20)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        count = count + 1

    p = document.add_paragraph()
    p.add_run("DEFINITIONS").underline = True
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    arrLen = len(definitionsArray)
    count_two = 0

    for defn in definitionsArray:
      if count_two == arrLen:
         break
      
      paragraph = document.add_paragraph(f"{count_two + 1}. {defn}")
      paragraph.paragraph_format.line_spacing = Pt(20)
      paragraph.paragraph_format.space_after = Pt(12)
      paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
      count_two = count_two + 1

    return document